#!/usr/bin/env python3
"""Fill a conference-abstract template .docx with metadata-driven content.

Reads a YAML config describing title, authors, body, figure, caption, and
writes a filled single-page docx. Designed for AAICon-style templates where
the top-matter is a 4-row title/author table and the figure sits in a
2-row caption table, but tolerates simple structural variation.

Usage:
    build_abstract.py --config config.yaml --out output.docx [--template path.docx]

Body markup:
    Use [b]text[/b] for bold spans and [url=https://...]text[/url] for hyperlinks.
    Newlines in body are treated as paragraph breaks.
"""
from __future__ import annotations

import argparse
import re
from pathlib import Path

import yaml
from docx import Document
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Mm, Pt, RGBColor

SKILL_ROOT = Path(__file__).resolve().parents[1]
DEFAULT_TEMPLATE = SKILL_ROOT / "assets" / "aaicon_template.docx"
GENERATE_SCRIPT = SKILL_ROOT / "scripts" / "generate_figure.py"


def _resolve_or_generate_figure(fig_cfg: dict, config_path, out_path) -> Path:
    """Return an existing figure image path, or generate one via Codex.

    Resolution order:
      1. `figure.image` exists on disk            → use it
      2. `figure.generate` block present          → run generate_figure.py,
         cache the result under the config's dir, return that path
      3. otherwise                                → FileNotFoundError
    """
    cfg_dir = Path(config_path).resolve().parent
    if fig_cfg.get("image"):
        img = Path(fig_cfg["image"]).expanduser()
        if not img.is_absolute():
            img = (cfg_dir / img).resolve()
        if img.exists():
            return img
        if not fig_cfg.get("generate"):
            raise FileNotFoundError(
                f"figure.image not found: {img}. "
                "Use an absolute path, relative-to-config path, "
                "or add a `figure.generate` block."
            )

    gen = fig_cfg.get("generate")
    if not gen:
        raise FileNotFoundError(
            "figure.image is missing and figure.generate is not set — "
            "cannot produce a figure."
        )

    # Default cached filename sits next to the config (config.dir / generated/…)
    cache_name = gen.get("file_name") or (Path(out_path).stem + "_figure.png")
    cache_dir = Path(gen.get("output_dir") or (cfg_dir / "generated")).expanduser()
    if not cache_dir.is_absolute():
        cache_dir = (cfg_dir / cache_dir).resolve()
    cache_path = cache_dir / cache_name

    if cache_path.exists() and not gen.get("force", False):
        print(f"[figure] reusing cached image: {cache_path}")
        return cache_path

    import subprocess, sys as _sys
    cmd = [
        _sys.executable, str(GENERATE_SCRIPT),
        "--prompt", gen["prompt"],
        "--out", str(cache_path),
        "--aspect", gen.get("aspect", "landscape"),
        "--background", gen.get("background", "auto"),
        "--timeout", str(gen.get("timeout_seconds", 180)),
    ]
    if gen.get("codex"):
        cmd += ["--codex", gen["codex"]]
    print(f"[figure] generating via Codex: {' '.join(cmd)}")
    r = subprocess.run(cmd)
    if r.returncode != 0 or not cache_path.exists():
        raise RuntimeError(
            f"figure generation failed (exit {r.returncode}). "
            f"See logs in {cache_dir / (cache_path.stem + '.codex_logs')}"
        )
    return cache_path

KOREAN_FONT = "맑은 고딕"
ASCII_FONT = "Malgun Gothic"

BOLD_RE = re.compile(r"\[b\](.+?)\[/b\]", re.DOTALL)
URL_RE = re.compile(r"\[url=([^\]]+)\]([^\[]+)\[/url\]")
TOKEN_RE = re.compile(r"(\[b\].+?\[/b\]|\[url=[^\]]+\][^\[]+\[/url\])", re.DOTALL)


def set_run_font(run, size_pt=None, bold=None,
                 font_ascii=ASCII_FONT, font_ea=KOREAN_FONT, color=None):
    if bold is not None:
        run.bold = bold
    if size_pt is not None:
        run.font.size = Pt(size_pt)
    if color is not None:
        run.font.color.rgb = RGBColor(*color)
    rPr = run._element.get_or_add_rPr()
    rFonts = rPr.find(qn("w:rFonts"))
    if rFonts is None:
        rFonts = OxmlElement("w:rFonts")
        rPr.insert(0, rFonts)
    rFonts.set(qn("w:ascii"), font_ascii)
    rFonts.set(qn("w:hAnsi"), font_ascii)
    rFonts.set(qn("w:eastAsia"), font_ea)
    rFonts.set(qn("w:cs"), font_ascii)


def clear_paragraph(p):
    """Remove all runs in a paragraph, keeping its pPr (formatting)."""
    for r in list(p.runs):
        r._element.getparent().remove(r._element)
    # Also kill hyperlinks
    for hl in p._element.findall(qn("w:hyperlink")):
        p._element.remove(hl)


def add_hyperlink(paragraph, url, text, size_pt=None, bold=False):
    """Insert a clickable hyperlink run into a paragraph."""
    part = paragraph.part
    r_id = part.relate_to(
        url,
        "http://schemas.openxmlformats.org/officeDocument/2006/relationships/hyperlink",
        is_external=True,
    )
    hl = OxmlElement("w:hyperlink")
    hl.set(qn("r:id"), r_id)
    new_run = OxmlElement("w:r")
    rPr = OxmlElement("w:rPr")
    # styled: blue + underline
    color = OxmlElement("w:color"); color.set(qn("w:val"), "0563C1"); rPr.append(color)
    u = OxmlElement("w:u"); u.set(qn("w:val"), "single"); rPr.append(u)
    if size_pt is not None:
        sz = OxmlElement("w:sz"); sz.set(qn("w:val"), str(int(size_pt * 2))); rPr.append(sz)
    if bold:
        b = OxmlElement("w:b"); rPr.append(b)
    rFonts = OxmlElement("w:rFonts")
    rFonts.set(qn("w:ascii"), ASCII_FONT)
    rFonts.set(qn("w:hAnsi"), ASCII_FONT)
    rFonts.set(qn("w:eastAsia"), KOREAN_FONT)
    rPr.append(rFonts)
    new_run.append(rPr)
    t = OxmlElement("w:t"); t.text = text; t.set(qn("xml:space"), "preserve")
    new_run.append(t)
    hl.append(new_run)
    paragraph._element.append(hl)


def add_inline_markup(paragraph, text, size_pt, default_bold=False):
    """Parse [b]...[/b] and [url=...]...[/url] markers and emit runs."""
    if not text:
        return
    pieces = TOKEN_RE.split(text)
    for piece in pieces:
        if not piece:
            continue
        m = BOLD_RE.fullmatch(piece)
        if m:
            r = paragraph.add_run(m.group(1))
            set_run_font(r, size_pt=size_pt, bold=True)
            continue
        m = URL_RE.fullmatch(piece)
        if m:
            add_hyperlink(paragraph, m.group(1), m.group(2), size_pt=size_pt)
            continue
        r = paragraph.add_run(piece)
        set_run_font(r, size_pt=size_pt, bold=default_bold)


def set_cell_text(cell, lines, size_pt, bold=False, alignment=None, line_spacing=1.15):
    """Replace cell contents with one paragraph per line entry.

    `lines` items may be a str or a dict {text, size_pt?, bold?}. Inline
    markup ([b], [url]) is expanded.
    """
    # Remove every paragraph in the cell except the first (which we reuse)
    first_p = cell.paragraphs[0]
    for p in cell.paragraphs[1:]:
        p._element.getparent().remove(p._element)
    clear_paragraph(first_p)

    def _apply(p, entry):
        if alignment is not None:
            p.alignment = alignment
        pf = p.paragraph_format
        pf.line_spacing = line_spacing
        pf.space_before = Pt(0)
        pf.space_after = Pt(0)
        if isinstance(entry, dict):
            text = entry.get("text", "")
            local_size = entry.get("size_pt", size_pt)
            local_bold = entry.get("bold", bold)
        else:
            text = entry
            local_size = size_pt
            local_bold = bold
        add_inline_markup(p, text, size_pt=local_size, default_bold=local_bold)

    _apply(first_p, lines[0])
    for extra in lines[1:]:
        p = cell.add_paragraph()
        _apply(p, extra)


def delete_table(table):
    table._element.getparent().remove(table._element)


def replace_body_block(doc, start_paragraph_text_contains, new_blocks,
                       body_size_pt=9, first_line_indent_mm=4,
                       justify=True, line_spacing=1.0):
    """Replace existing body paragraphs after the heading with new blocks.

    Finds the first paragraph whose text contains the given substring
    (e.g. '요   약'), deletes every following paragraph up to (but not
    including) the next table in the body, then inserts the new body blocks
    into that gap.

    `new_blocks` is a list of strings. Each string becomes one paragraph.
    Inline markup ([b], [url]) is expanded.
    """
    body = doc.element.body
    children = list(body.iterchildren())
    # find heading index
    heading_idx = None
    for i, child in enumerate(children):
        if child.tag.endswith("}p") and start_paragraph_text_contains in "".join(
            t.text or "" for t in child.iter(qn("w:t"))
        ):
            heading_idx = i
            break
    if heading_idx is None:
        raise RuntimeError(f"Could not find heading containing '{start_paragraph_text_contains}'")

    # Delete everything after heading until the next tbl element
    i = heading_idx + 1
    anchor = None
    while i < len(children):
        child = children[i]
        if child.tag.endswith("}tbl"):
            anchor = child
            break
        body.remove(child)
        children.pop(i)
    if anchor is None:
        raise RuntimeError("No table found after the abstract heading (expected figure table).")

    # Create paragraph XML template from heading to inherit style, then customize
    for block in new_blocks:
        p = OxmlElement("w:p")
        pPr = OxmlElement("w:pPr")
        if justify:
            jc = OxmlElement("w:jc"); jc.set(qn("w:val"), "both"); pPr.append(jc)
        if first_line_indent_mm:
            ind = OxmlElement("w:ind")
            ind.set(qn("w:firstLine"), str(int(first_line_indent_mm * 56.7)))  # ~twips
            pPr.append(ind)
        spacing = OxmlElement("w:spacing")
        spacing.set(qn("w:line"), str(int(line_spacing * 240)))
        spacing.set(qn("w:lineRule"), "auto")
        spacing.set(qn("w:before"), "0")
        spacing.set(qn("w:after"), "0")
        pPr.append(spacing)
        p.append(pPr)
        anchor.addprevious(p)
    # Fill the freshly inserted (empty) paragraphs with styled runs
    paragraphs = doc.paragraphs
    # Find index of heading in doc.paragraphs
    heading_p_idx = next(
        i for i, p in enumerate(paragraphs)
        if start_paragraph_text_contains in p.text
    )
    # Now empty paragraphs inserted follow heading_p_idx+1 .. matching len(new_blocks)
    for offset, block in enumerate(new_blocks, start=1):
        p = paragraphs[heading_p_idx + offset]
        add_inline_markup(p, block, size_pt=body_size_pt)


def fill_figure_table(doc, figure_table_index, image_path, caption,
                      image_width_mm=140, caption_size_pt=8, caption_bold_label="그림 1."):
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    table = doc.tables[figure_table_index]
    # Image row
    img_cell = table.rows[0].cells[0]
    # clear paragraphs in cell, keep first
    first_p = img_cell.paragraphs[0]
    for p in img_cell.paragraphs[1:]:
        p._element.getparent().remove(p._element)
    clear_paragraph(first_p)
    first_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = first_p.add_run()
    run.add_picture(str(image_path), width=Mm(image_width_mm))

    # Caption row
    cap_cell = table.rows[1].cells[0]
    first_p = cap_cell.paragraphs[0]
    for p in cap_cell.paragraphs[1:]:
        p._element.getparent().remove(p._element)
    clear_paragraph(first_p)
    first_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    if caption_bold_label:
        r = first_p.add_run(caption_bold_label + " ")
        set_run_font(r, size_pt=caption_size_pt, bold=True)
    add_inline_markup(first_p, caption, size_pt=caption_size_pt)


def build(config_path: Path, out_path: Path, template_path: Path):
    from docx.enum.text import WD_ALIGN_PARAGRAPH

    cfg = yaml.safe_load(Path(config_path).read_text(encoding="utf-8"))
    doc = Document(str(template_path))

    # ---- Fill title/author table (Table 0) ----
    title_table = doc.tables[0]

    # Row 0: [tag] + KO title
    tag = cfg.get("tag", "구두")  # 구두 or 포스터
    ko_title_lines = cfg["title_ko"].splitlines() or [cfg["title_ko"]]
    set_cell_text(
        title_table.cell(0, 0),
        [f"[{tag}]"] + ko_title_lines,
        size_pt=16,
        bold=True,
        alignment=WD_ALIGN_PARAGRAPH.CENTER,
        line_spacing=1.15,
    )
    # First line [tag] should NOT be bold/16pt — rewrite first paragraph as 10pt normal right-aligned look? Template puts it inline with small font. Keep simple: render as 10pt normal top line.
    first_p = title_table.cell(0, 0).paragraphs[0]
    clear_paragraph(first_p)
    first_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = first_p.add_run(f"[{tag}]")
    set_run_font(r, size_pt=10, bold=False)

    # Row 1: KO author info
    ko_authors = cfg["authors_ko"]  # list of strings, each a line
    set_cell_text(
        title_table.cell(1, 0),
        ko_authors,
        size_pt=10,
        alignment=WD_ALIGN_PARAGRAPH.CENTER,
        line_spacing=1.15,
    )

    # Row 2: EN title
    en_title_lines = cfg["title_en"].splitlines()
    set_cell_text(
        title_table.cell(2, 0),
        en_title_lines,
        size_pt=14,
        alignment=WD_ALIGN_PARAGRAPH.CENTER,
        line_spacing=1.15,
    )

    # Row 3: EN author info
    set_cell_text(
        title_table.cell(3, 0),
        cfg["authors_en"],
        size_pt=10,
        alignment=WD_ALIGN_PARAGRAPH.CENTER,
        line_spacing=1.15,
    )

    # ---- Replace body ----
    body_blocks = cfg["body"] if isinstance(cfg["body"], list) else [cfg["body"]]
    replace_body_block(
        doc,
        start_paragraph_text_contains=cfg.get("abstract_heading", "요   약"),
        new_blocks=body_blocks,
        body_size_pt=cfg.get("body_size_pt", 9),
        first_line_indent_mm=cfg.get("first_line_indent_mm", 4),
        justify=cfg.get("justify", True),
        line_spacing=cfg.get("body_line_spacing", 1.0),
    )

    # ---- Fill figure table ----
    fig_cfg = cfg.get("figure")
    if fig_cfg:
        img_path = _resolve_or_generate_figure(fig_cfg, config_path, out_path)
        fill_figure_table(
            doc,
            figure_table_index=cfg.get("figure_table_index", 1),
            image_path=img_path,
            caption=fig_cfg["caption"],
            image_width_mm=fig_cfg.get("width_mm", 140),
            caption_size_pt=fig_cfg.get("caption_size_pt", 8),
            caption_bold_label=fig_cfg.get("caption_label", "그림 1."),
        )

    # ---- Remove guidelines table if requested ----
    if cfg.get("remove_guidelines", True):
        # Guidelines table is the last (third) table in the default template
        guide_idx = cfg.get("guidelines_table_index", 2)
        if guide_idx < len(doc.tables):
            delete_table(doc.tables[guide_idx])

    out_path.parent.mkdir(parents=True, exist_ok=True)
    doc.save(str(out_path))
    print(f"Saved: {out_path}  ({out_path.stat().st_size:,} bytes)")


def main():
    ap = argparse.ArgumentParser()
    ap.add_argument("--config", required=True, help="YAML config path")
    ap.add_argument("--out", required=True, help="Output .docx path")
    ap.add_argument("--template", default=str(DEFAULT_TEMPLATE),
                    help="Template .docx path (default: bundled AAICon template)")
    args = ap.parse_args()
    build(Path(args.config), Path(args.out), Path(args.template))


if __name__ == "__main__":
    main()
